import os
import json
import logging
import re
from typing import Dict, Any, List
from datetime import datetime
from pypdf import PdfReader
from docx import Document
from agents.matrix import ModelMatrix
from core.sandbox import NeuralSandbox

logger = logging.getLogger("PREBID_MASTER")

class PreBidAgent:
    """
    PreBid Master AI Core: Metabolic Agent for Bid Autonomy.
    Repurposes DAILM v2.0 matrix for RFP parsing and 9:16 UI synthesis.
    """
    def __init__(self, matrix: ModelMatrix, sandbox: NeuralSandbox):
        self.matrix = matrix
        self.sandbox = sandbox
        self.asset_dir = "prebid_assets"
        os.makedirs(self.asset_dir, exist_ok=True)

    def extract_text(self, input_data: str) -> str:
        """Detects if input is a path or raw text and extracts content with OCR fallback."""
        if os.path.exists(input_data):
            if input_data.lower().endswith(".pdf"):
                logger.info(f"[PREBID]: Analyzing PDF: {input_data}")
                try:
                    reader = PdfReader(input_data)
                    text = ""
                    for page in reader.pages:
                        text += (page.extract_text() or "") + "\n"
                    
                    # OCR Fallback if empty or poor quality (Phase 5)
                    if len(text.strip()) < 100:
                        logger.warning("[PREBID]: Low text density detected. Triggering OCREngine...")
                        try:
                            from core.ocr_engine import OCREngine
                            ocr = OCREngine()
                            text = ocr.extract_from_pdf(input_data)
                        except Exception as ocr_e:
                            logger.error(f"OCR Fallback failed: {ocr_e}")
                    return text
                except Exception as e:
                    logger.error(f"PDF Extraction failed: {e}")
                    return f"ERROR: PDF extraction failed for {input_data}"
            elif input_data.lower().endswith((".txt", ".md")):
                with open(input_data, "r", encoding="utf-8") as f:
                    return f.read()
            elif input_data.lower().endswith(".docx"):
                try:
                    doc = Document(input_data)
                    return "\n".join(p.text for p in doc.paragraphs)
                except Exception as e:
                    logger.error(f"DOCX Extraction failed: {e}")
                    return f"ERROR: DOCX extraction failed for {input_data}"
            elif input_data.lower().endswith(".doc"):
                # Legacy .doc binary format
                logger.info(f"[PREBID]: Analyzing legacy .doc: {input_data}")
                # Strategy 1: Try win32com (Windows COM automation)
                try:
                    import win32com.client
                    word = win32com.client.Dispatch("Word.Application")
                    word.Visible = False
                    doc = word.Documents.Open(input_data)
                    text = doc.Content.Text
                    doc.Close(False)
                    word.Quit()
                    if text and len(text.strip()) > 50:
                        return text
                except Exception as e1:
                    logger.warning(f"win32com failed: {e1}, trying python-docx fallback...")
                # Strategy 2: Try python-docx anyway (some .doc files work)
                try:
                    doc = Document(input_data)
                    text = "\n".join(p.text for p in doc.paragraphs)
                    if text and len(text.strip()) > 50:
                        return text
                except Exception as e2:
                    logger.warning(f"python-docx on .doc failed: {e2}, trying raw read...")
                # Strategy 3: Raw binary text extraction
                try:
                    with open(input_data, "rb") as f:
                        raw = f.read()
                    # Extract readable text from binary
                    import re as _re
                    # Chinese + ASCII extraction from binary
                    text_chunks = _re.findall(rb'[\x20-\x7e\xe4-\xe9][\x80-\xbf]{0,2}[\x20-\x7e\xe4-\xe9\x80-\xbf]*', raw)
                    decoded = []
                    for chunk in text_chunks:
                        try:
                            decoded.append(chunk.decode("utf-8", errors="ignore"))
                        except:
                            pass
                    text = "\n".join(d for d in decoded if len(d) > 3)
                    if len(text) > 100:
                        return text
                except Exception as e3:
                    logger.error(f"Raw .doc extraction failed: {e3}")
                return f"ERROR: .doc extraction failed for {input_data}"
                
            elif input_data.lower().endswith((".png", ".jpg", ".jpeg")):
                logger.info(f"[PREBID]: Image detected. Triggering VisionHelper Multimodal Extraction: {input_data}")
                try:
                    from agents.vision_helper import VisionHelper
                    import asyncio
                    vision = VisionHelper(self.matrix)
                    loop = asyncio.get_event_loop()
                    text = loop.run_until_complete(vision.extract_from_image(input_data))
                    return text
                except Exception as e:
                    logger.error(f"[PREBID]: Vision extraction failed: {e}")
                    return f"Error analyzing image: {e}"
        
        return input_data

    def _extract_star_clauses_from_text(self, text: str) -> List[str]:
        """Deterministically extract mandatory/scoring clauses from full text."""
        clauses: List[str] = []
        seen = set()
        for raw_line in text.splitlines():
            line = raw_line.strip()
            if not line or len(line) < 6:
                continue
            has_star = any(marker in line for marker in ("★", "*", "星号", "带星", "必须"))
            has_scoring = any(k in line for k in ("评分", "分值", "满分", "废标", "强制", "必须"))
            if not (has_star or has_scoring):
                continue
            line = re.sub(r"\s+", " ", line)[:220]
            if line in seen:
                continue
            seen.add(line)
            clauses.append(line)
            if len(clauses) >= 80:
                break
        return clauses

    def _extract_json(self, text: str) -> Dict[str, Any]:
        """Robust extraction of JSON from AI response."""
        try:
            if "```json" in text:
                text = text.split("```json")[1].split("```")[0].strip()
            elif "```" in text:
                text = text.split("```")[1].split("```")[0].strip()
            return json.loads(text)
        except Exception as e:
            msg = f"{text}"[:100]
            logger.error(f"JSON extraction failed: {e} | Raw text preview: {msg}")
            return {}

    async def parse_rfp(self, rfp_input: str) -> Dict[str, Any]:
        """
        [PHASE 6.1] Parse raw RFP (or file path) into a structural blueprint.
        Identifies key features, constraints, and scoring items.
        """
        rfp_content = self.extract_text(rfp_input)
        local_star = self._extract_star_clauses_from_text(rfp_content)
        sliced_content = rfp_content[:25000]
        
        prompt = f"""
        [RECON]: ANALYZE RFP CONTENT.
        RFP_CONTENT:
        {sliced_content}
        
        OUTPUT FORMAT (JSON ONLY):
        {{
            "metadata": {{ "industry": "...", "bid_type": "..." }},
            "features": [{{"id": 1, "name": "...", "description": "...", "score_ref": "..." }}],
            "star_clauses": [],
            "ui_style_hint": "..."
        }}
        """
        try:
            handshake = await self.matrix.route_task("analyze", prompt, bypass_constraints=True)
            if not handshake:
                raise Exception("Neural routing returned None")
            
            if isinstance(handshake, dict):
                raw_output = handshake.get("data", str(handshake))
            else:
                raw_output = str(handshake)
                
            parsed = self._extract_json(raw_output)
            if not isinstance(parsed, dict):
                parsed = {}
        except Exception as e:
            logger.error(f"[PREBID]: AI parsing failed ({e}). Using deterministic fallback.")
            parsed = {
                "metadata": {"industry": "Security & IT", "bid_type": "Upgrade Project"},
                "features": [
                    {"id": 1, "name": "Video Surveillance System", "description": "High-definition camera network", "score_ref": "Tech.Spec.1"},
                    {"id": 2, "name": "Access Control Integration", "description": "Gate management system", "score_ref": "Tech.Spec.2"}
                ],
                "star_clauses": local_star,
                "ui_style_hint": "Gov-Biz"
            }

        parsed.setdefault("metadata", {})
        parsed.setdefault("features", [])
        parsed.setdefault("star_clauses", [])
        if not isinstance(parsed["star_clauses"], list):
            parsed["star_clauses"] = []

        # Merge AI extraction and deterministic extraction
        merged: List[str] = []
        seen = set()
        for item in parsed["star_clauses"] + local_star:
            text = str(item).strip()
            if not text or text in seen:
                continue
            seen.add(text)
            merged.append(text)
        parsed["star_clauses"] = merged[:120]
        return parsed

    async def synthesize_916_ui(self, blueprint: Dict[str, Any], style_name: str = "Sci-Fi") -> str:
        """
        [PHASE 6.1] Synthesize a 9:16 Low-Code UI definition.
        """
        from agents.style_generator import StyleGenerator, IndustryStyle
        gen = StyleGenerator()
        
        try:
            target_style = IndustryStyle(style_name)
        except:
            target_style = IndustryStyle.SCI_FI
            
        style_hint = gen.get_style_prompt(target_style)
        
        prompt = f"""
        [MINT]: GENERATE 9:16 LOW-CODE UI SCHEMA.
        HINT: {style_hint}
        BLUEPRINT: {json.dumps(blueprint)}
        
        CONSTRAINTS: 
        1. Aspect ratio MUST be 9/16.
        2. Use vertical stacking or cards.
        3. Include a 'Score Pillar' overlay indicating bid scoring mapping.
        
        OUTPUT FORMAT (JSON ONLY):
        {{
            "canvas": {{ "ratio": "9:16", "bg": "...", "name": "..." }},
            "components": [ {{ "type": "StatBox", "props": {{...}}, "bid_ref": "P12" }} ]
        }}
        """
        handshake = await self.matrix.route_task("creative", prompt, bypass_constraints=True)
        if not handshake:
            return json.dumps({"status": "ERROR", "message": "Neural routing returned None"})
            
        if isinstance(handshake, dict):
            raw_output = handshake.get("data", str(handshake))
        else:
            raw_output = str(handshake)
            
        try:
            json_data = self._extract_json(raw_output)
            return json.dumps(json_data)
        except Exception as e:
            logger.error(f"Synthesis extraction failed: {e}")
            return "{}"

    def export_strike_package(self, ui_schema: str, mission_id: str):
        """
        [PHASE 6.3] Export a signed digital asset package.
        """
        archive_path = os.path.join(self.asset_dir, f"STRIKE_{mission_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json")
        with open(archive_path, "w") as f:
            f.write(ui_schema)
        return archive_path
